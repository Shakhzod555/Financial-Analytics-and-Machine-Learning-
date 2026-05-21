"""
Build a self-contained Word report from the replication outputs.
Produces report.docx with body text + tables + figures + references.
"""

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(".")
DOC = Document()


# ---------------------------------------------------------------------
# Global styling
# ---------------------------------------------------------------------
def set_global_styles(doc):
    # Normal body
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)


def add_para(doc, text, *, style=None, align=None, bold=False, italic=False,
             size=None, space_after=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    return p


def add_rich_para(doc, runs, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_line_indent=None, space_after=4):
    """Add a paragraph from a list of (text, **kwargs) tuples."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    for item in runs:
        if isinstance(item, str):
            text, opts = item, {}
        else:
            text, opts = item
        r = p.add_run(text)
        r.bold = opts.get("bold", False)
        r.italic = opts.get("italic", False)
        if "size" in opts:
            r.font.size = Pt(opts["size"])
        if "color" in opts:
            r.font.color.rgb = opts["color"]
    return p


def add_heading(doc, text, level=1, *, color=None):
    h = doc.add_heading(text, level=level)
    if color is not None:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def add_page_break(doc):
    doc.add_page_break()


def add_caption(doc, label, caption_text, *, italic_label=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    r1 = p.add_run(label + " ")
    r1.bold = True
    if italic_label:
        r1.italic = False
    r2 = p.add_run(caption_text)
    r2.font.size = Pt(10)


def add_image(doc, path, width_in=6.5, *, label=None, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(path, width=Inches(width_in))
    if label and caption:
        add_caption(doc, label, caption)


# Table formatting helpers
def _set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_cell_border(cell, *, top=None, bottom=None, left=None, right=None):
    """Apply only specified borders; pass dict {sz: '8', val: 'single', color: '000000'}."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, spec in [("top", top), ("bottom", bottom),
                            ("left", left), ("right", right)]:
        existing = tc_borders.find(qn(f"w:{edge_name}"))
        if existing is not None:
            tc_borders.remove(existing)
        if spec is None:
            edge = OxmlElement(f"w:{edge_name}")
            edge.set(qn("w:val"), "nil")
        else:
            edge = OxmlElement(f"w:{edge_name}")
            edge.set(qn("w:val"), spec.get("val", "single"))
            edge.set(qn("w:sz"), spec.get("sz", "6"))
            edge.set(qn("w:space"), "0")
            edge.set(qn("w:color"), spec.get("color", "000000"))
        tc_borders.append(edge)


def make_dataframe_table(doc, df, *, font_size=9, header_bold=True,
                         booktabs=True, col_widths=None):
    """Render a DataFrame as a docx table with booktabs-like horizontal rules."""
    n_rows, n_cols = df.shape[0] + 1, df.shape[1]
    tab = doc.add_table(rows=n_rows, cols=n_cols)
    tab.alignment = WD_TABLE_ALIGNMENT.CENTER
    tab.autofit = True

    # Header row
    for j, col in enumerate(df.columns):
        cell = tab.cell(0, j)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(col))
        run.bold = header_bold
        run.font.size = Pt(font_size)

    # Data rows
    for i in range(df.shape[0]):
        for j in range(df.shape[1]):
            cell = tab.cell(i + 1, j)
            cell.text = ""
            p = cell.paragraphs[0]
            val = df.iat[i, j]
            text = "" if pd.isna(val) else str(val)
            run = p.add_run(text)
            run.font.size = Pt(font_size)
            # Right-align numeric-looking cells (after column 1 typically)
            if j >= 1 and any(ch.isdigit() for ch in text):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Apply booktabs-style borders (clear all, then top/bottom of header + bottom of last row)
    if booktabs:
        for row in tab.rows:
            for cell in row.cells:
                _set_cell_border(cell, top=None, bottom=None, left=None, right=None)
        # Top rule above header
        for cell in tab.rows[0].cells:
            _set_cell_border(cell, top={"sz": "12"}, bottom={"sz": "6"})
        # Bottom rule under last row
        for cell in tab.rows[-1].cells:
            _set_cell_border(cell, bottom={"sz": "12"})

    return tab


# ---------------------------------------------------------------------
# Content
# ---------------------------------------------------------------------
def build_report():
    set_global_styles(DOC)
    nav = WD_ALIGN_PARAGRAPH

    # ===== TITLE BLOCK =====
    title = DOC.add_paragraph()
    title.alignment = nav.CENTER
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run("A Machine Learning Approach to Volatility Forecasting")
    r.bold = True
    r.font.size = Pt(15)

    sub = DOC.add_paragraph()
    sub.alignment = nav.CENTER
    sub.paragraph_format.space_after = Pt(2)
    r = sub.add_run("Replication of Christensen, Siggaard & Veliyev "
                    "(2023, Journal of Financial Econometrics)")
    r.italic = True
    r.font.size = Pt(11)

    aff = DOC.add_paragraph()
    aff.alignment = nav.CENTER
    aff.paragraph_format.space_after = Pt(14)
    r = aff.add_run("Financial Analytics and Machine Learning — Coursework Report")
    r.font.size = Pt(10)

    # ===== SECTION 1 — THE PAPER =====
    add_heading(DOC, "1. The paper", level=1)

    add_rich_para(DOC, [
        ("Research question. ", {"bold": True}),
        "Forecasting volatility underpins risk management (VaR), option "
        "pricing, and dynamic portfolio allocation, so even small gains "
        "in forecast accuracy translate into measurable economic value. "
        "Realised variance (RV) — the sum of squared intraday returns — "
        "is the dominant proxy for latent volatility (Andersen & "
        "Bollerslev, 1998). Corsi's (2009) Heterogeneous Autoregressive "
        "(HAR) model — daily RV regressed on its 1-, 5- and 22-day lags "
        "— captures the long-memory of volatility with three parameters "
        "and remains the benchmark in the literature. Christensen, "
        "Siggaard & Veliyev (2023, henceforth CSV) ask whether modern "
        "machine-learning techniques can extract additional predictive "
        "content either from the same three HAR lags or from a richer "
        "cross-section of volatility correlates, and they ask why any "
        "improvement arises.",
    ])

    add_rich_para(DOC, [
        ("Methodology. ", {"bold": True}),
        "CSV run a horse-race across 22 models on 29 DJIA stocks, "
        "2001–2017 (T = 4,257 sessions). The HAR lineage (HAR, LogHAR, "
        "LevHAR, SHAR, HARQ) is compared with three regularised linear "
        "models (Ridge, Lasso, Elastic Net), three tree ensembles "
        "(Bagging, Random Forest, Gradient Boosting), and four "
        "feed-forward neural networks of pyramid-shaped depth "
        "(NN₁–NN₄, both single-model and 10-of-100 ensemble variants). "
        "Two information sets are studied: M",
        ("HAR", {"italic": True}),
        " contains the three lagged-RV terms — used for a one-to-one "
        "comparison with HAR — while M",
        ("ALL", {"italic": True}),
        " adds nine covariates: option-implied volatility (IV), an "
        "earnings-announcement dummy, VIX, the Hang Seng squared return, "
        "the Aruoba–Diebold–Scotti business-conditions index, US 3-month "
        "T-bill changes, economic policy uncertainty (Baker et al. 2016), "
        "one-week momentum and dollar trading volume. Models are trained "
        "on 70 % of the sample, validated on 10 %, and evaluated on the "
        "final 20 %. Predictive accuracy is compared with the "
        "Diebold–Mariano (1995) test and the Model Confidence Set of "
        "Hansen, Lunde & Nason (2011); marginal effects are inspected "
        "with Accumulated Local Effect plots (Apley & Zhu, 2020).",
    ])

    add_rich_para(DOC, [
        ("Findings. ", {"bold": True}),
        "The paper makes five claims. ",
        ("(i)", {"bold": True}),
        " Off-the-shelf ML beats HAR even on M",
        ("HAR", {"italic": True}),
        ", with the differences DM-significant on more than half of stocks "
        "for the neural network. ",
        ("(ii)", {"bold": True}),
        " On M",
        ("ALL", {"italic": True}),
        ", regularisation and tree ensembles deliver an 8–10 % MSE "
        "reduction relative to HAR, while the unregularised extended HAR "
        "overfits. ",
        ("(iii)", {"bold": True}),
        " ALE plots show agreement on the most important predictors "
        "(RVD, RVW, IV) but disagreement on their ranking. ",
        ("(iv)", {"bold": True}),
        " Forecast gains rise with horizon — at h = 22 days the Random "
        "Forest reduces MSE by ~40 % vs HAR — and the authors attribute "
        "this to ML's ability to approximate long-memory persistence, "
        "documented by a slower autocorrelation decay of the in-sample "
        "fitted RV series (their Figure 8). ",
        ("(v)", {"bold": True}),
        " A Value-at-Risk application confirms the qualitative picture "
        "with smaller but still favourable ML gains.",
    ])

    add_rich_para(DOC, [
        ("Replication scope. ", {"bold": True}),
        "We target findings (i), (iv) and the ACF mechanism behind (iv). "
        "Reconstructing M",
        ("ALL", {"italic": True}),
        " is infeasible here: option-implied volatility, EPU, ADS and the "
        "earnings calendar are not in the available dataset. We therefore "
        "focus on the cleanest comparison — HAR vs ML on M",
        ("HAR", {"italic": True}),
        " — and on the ACF mechanism, which is independent of the "
        "information set and is the paper's main explanatory claim "
        "rather than a head-line result.",
    ])

    add_page_break(DOC)

    # ===== SECTION 2 — IMPLEMENTATION & HEADLINE RESULTS =====
    add_heading(DOC, "2. Implementation and headline results", level=1)

    add_rich_para(DOC, [
        ("Data and RV construction. ", {"bold": True}),
        "Our sample is 1-minute OHLCV for AAPL (Technology), AMZN "
        "(Consumer Discretionary) and JPM (Financials), spanning "
        "4 January 2016 to 31 December 2024, 2,264 NYSE sessions of "
        "390 bars per day. Following Andersen & Bollerslev (1998), 1-minute "
        "closes are subsampled to 5 minutes, log-returns are taken within "
        "each day (overnight returns excluded), and daily RV is computed "
        "as the sum of squared 5-minute returns. We additionally compute "
        "realised quarticity RQ, signed semivariance (RV⁺, RV⁻) and "
        "aggregated negative returns r⁻ to support LevHAR, SHAR and HARQ "
        "(Patton & Sheppard, 2015; Bollerslev, Patton & Quaedvlieg, 2016). "
        "Descriptive statistics are reported in Table A1; the unconditional "
        "RV path is plotted in Figure A1 and shows the COVID-19 spike "
        "(March 2020) and the August 2024 carry-trade unwind clearly.",
    ])

    add_rich_para(DOC, [
        ("Models. ", {"bold": True}),
        "Eight models are estimated. The HAR lineage is implemented with "
        "ordinary least squares: HAR (three lagged-RV averages), LogHAR "
        "(log-RV regression with Jensen bias correction "
        "exp(ŷ + 0.5σ̂²)), LevHAR (HAR augmented with three aggregated "
        "negative-return regressors), SHAR (RV decomposed into positive "
        "and negative semivariance) and HARQ (RVD interacted with √RQ "
        "to correct the measurement-error attenuation). The machine-"
        "learning side uses Elastic Net with hyperparameters chosen by "
        "grid search over α ∈ [0.1, 0.9] and λ ∈ [10⁻⁵, 10²] on the "
        "held-out 10 % validation set (CSV, Appendix A.4; k-fold CV is "
        "avoided to prevent autocorrelation-induced leakage), a Random "
        "Forest of 500 trees with min-leaf 5 and √p random features "
        "per split, and a feed-forward neural network with two hidden "
        "layers of 4 and 2 ReLU neurons trained with Adam at learning "
        "rate 0.001. Each NN "
        "is an ensemble of the top-3 of 10 random seeds ranked by "
        "validation MSE. ML features are standardised on the training "
        "sample; the NN's target is also standardised and "
        "inverse-transformed on prediction. Negative forecasts are "
        "floored at the in-sample RV minimum, following the paper's "
        "insanity filter.",
    ])

    add_rich_para(DOC, [
        ("Forecast targets and evaluation. ", {"bold": True}),
        "Three horizons are evaluated: h ∈ {1, 5, 22} trading days. "
        "The target at index t is the average future RV "
        "ȳₜ = h⁻¹ Σ RVₜ₊ₖ over k = 1…h. The 70/10/20 split places the "
        "test window at March 2023 – December 2024. We report "
        "out-of-sample mean squared error relative to HAR and conduct "
        "Diebold–Mariano (1995) tests on squared-error differentials "
        "with the Harvey, Leybourne & Newbold (1997) small-sample "
        "correction and a Newey–West-style long-run variance using "
        "h − 1 Bartlett lags. Main numerical results are in Table A2; "
        "Figure A2 visualises them.",
    ])

    add_rich_para(DOC, [
        ("Headline results. ", {"bold": True}),
        "Four patterns are clear. ",
        ("First", {"italic": True}),
        ", at h = 1 day the LogHAR model is the strongest in the HAR "
        "family on cross-section average (relative MSE 0.931, "
        "DM-significant at 1 % for JPM and 10 % for AAPL); the neural "
        "network is competitive on AAPL (0.928*), consistent with the "
        "paper's claim that even a shallow NN beats HAR on M",
        ("HAR", {"italic": True}),
        ". ",
        ("Second", {"italic": True}),
        ", LogHAR's relative advantage widens with horizon: rel-MSE "
        "0.931 → 0.837 → 0.877 across h = 1, 5, 22 days. ",
        ("Third", {"italic": True}),
        ", the Random Forest deteriorates with horizon on cross-section "
        "average — its monthly rel-MSE averages 2.142, driven by AMZN "
        "(3.11) and JPM (2.11) — but the picture is heterogeneous: on "
        "JPM at the weekly horizon RF in fact ",
        ("beats", {"italic": True}),
        " HAR significantly (0.835**), so the headline "
        "\"ML deteriorates with horizon\" hides a stock-level pattern "
        "that motivates Section 3. ",
        ("Fourth", {"italic": True}),
        ", Elastic Net offers no improvement on HAR under M",
        ("HAR", {"italic": True}),
        " (mean rel-MSE 1.05 – 1.12 across horizons, worst on JPM at "
        "1.39 at h = 5), consistent with the paper's argument that ML "
        "gains chiefly arise under the richer M",
        ("ALL", {"italic": True}),
        " information set.",
    ])

    add_page_break(DOC)

    # ===== SECTION 3 — DISCUSSION =====
    add_heading(DOC, "3. Critical comparison and discussion", level=1)

    add_rich_para(DOC, [
        ("What replicates. ", {"bold": True}),
        "Three of the paper's findings survive on our smaller and "
        "later sample. ",
        ("(a)", {"bold": True}),
        " At h = 1 day the HAR lineage is broadly indistinguishable "
        "except in the right tail, where LogHAR's log transform "
        "down-weights outliers: Figure A5 shows LogHAR's MSE drops to "
        "0.69 of HAR's in the highest RV decile for both AAPL and JPM. ",
        ("(b)", {"bold": True}),
        " The neural network beats HAR on AAPL at h = 1 with DM "
        "significance (0.928*), and again at h = 22 (0.826**), supporting "
        "the paper's claim that a two-layer (4, 2) network is sufficient "
        "to capture the relevant non-linearity. ",
        ("(c)", {"bold": True}),
        " The mechanism the paper proposes — that ML's stronger "
        "long-memory approximation drives long-horizon improvement — "
        "is independently visible in our sample. Figure A3 (our "
        "reproduction of CSV's Figure 8) shows that at h = 22 the "
        "autocorrelation of fitted RV decays substantially more slowly "
        "for RF and NN than for HAR. For AAPL at lag 60, HAR's ACF is "
        "below 0.10 while RF's is around 0.22 and NN's around 0.15; "
        "the same gap is visible for JPM and (less starkly) for AMZN.",
    ])

    add_rich_para(DOC, [
        ("What does not replicate. ", {"bold": True}),
        "The paper's central head-line — that ",
        ("MSE gains grow with horizon", {"italic": True}),
        " — is inverted in our sample. Random-Forest rel-MSE rises "
        "from 1.08 at h = 1 to 2.14 at h = 22 on average. We identify "
        "four candidate explanations. ",
        ("(1) Test-period regime. ", {"bold": True}),
        "The paper's evaluation window contains the 2008 Global "
        "Financial Crisis, the 2010 flash crash and 2011 European debt "
        "stress — high-vol shocks where ML's long-memory persistence "
        "yields large gains relative to a mean-reverting HAR. Our test "
        "window (March 2023 – December 2024) is the calmest in our "
        "sample (Figure A1). ",
        ("(2) Cross-section. ", {"bold": True}),
        "Averaging across 29 stocks washes out idiosyncratic blow-ups; "
        "with three stocks a single bad trajectory (AMZN) drives the "
        "average. ",
        ("(3) Information set. ", {"bold": True}),
        "CSV's most pronounced ML gains arise under M",
        ("ALL", {"italic": True}),
        ", where extra covariates discriminate ML from HAR; M",
        ("HAR", {"italic": True}),
        " offers ML no comparative advantage. ",
        ("(4) Rolling-window refit. ", {"bold": True}),
        "The paper re-estimates day-by-day. Re-running HAR and RF with "
        "a 22-day expanding-window refit leaves the picture essentially "
        "unchanged (Figure A6: AMZN 3.11 → 2.70, JPM 2.11, AAPL 1.20 → "
        "1.18); rolling refit is not the missing ingredient and "
        "hypothesis (1) — the calm-regime mismatch — remains the "
        "leading explanation.",
    ])

    add_rich_para(DOC, [
        ("Loss-function robustness. ", {"bold": True}),
        "MSE on realised variance is sensitive to proxy noise (Patton, "
        "2011); the QLIKE loss σ²/σ̂² − log(σ²/σ̂²) − 1 preserves the "
        "consistency of the optimal forecast under that noise. Table A3 "
        "re-evaluates the contest under QLIKE. The ranking is preserved "
        "— LogHAR remains best at h = 22 (rel-QLIKE 0.823*) — but the "
        "magnitudes change. RF's monthly deterioration shrinks from "
        "2.142 (MSE) to 1.296 (QLIKE); on AAPL at h = 22, RF in fact ",
        ("beats", {"italic": True}),
        " HAR under QLIKE (0.873) where it lost under MSE (1.200). HARQ "
        "behaves similarly: rel-MSE 1.190 → rel-QLIKE 0.959, consistent "
        "with HARQ amplifying noisy quarticity into large ",
        ("absolute", {"italic": True}),
        " errors that QLIKE discounts. The inverted long-horizon ML "
        "result is therefore partly a loss-function artefact, though "
        "the AMZN deterioration survives QLIKE (1.885) and the regime "
        "mismatch above remains the dominant explanation. For "
        "risk-management applications where the relative mispricing of "
        "variance matters more than its absolute level — e.g., VaR "
        "scaling or option-implied surfaces — the QLIKE-based ranking "
        "is the operationally relevant one.",
    ])

    add_rich_para(DOC, [
        ("Insight. ", {"bold": True}),
        "The replication produces an ",
        ("interpretable", {"italic": True}),
        " divergence. The structural property the paper identifies — "
        "ML's long-memory persistence — is robust across our smaller "
        "cross-section (Figure A3) even though it does not deliver an "
        "MSE gain. This decoupling is informative: the persistence is an "
        "unconditional property of the fitted models, while its "
        "MSE-payoff is conditional on the realised volatility process "
        "during the evaluation window exhibiting the persistent shocks "
        "the models are pre-disposed to forecast. In the calm post-COVID "
        "test period, that property turns into a liability — RF and NN "
        "over-predict mean-reversion from the COVID era into a quiescent "
        "2023–24, with the damage concentrated in the middle RV deciles "
        "(Figure A5) that dominate our test window. The paper's "
        "long-horizon claim therefore appears ",
        ("regime-conditional rather than universal", {"italic": True}),
        ", and partly ",
        ("loss-function-conditional", {"italic": True}),
        " — nuances the paper's wide cross-section is large enough to "
        "obscure. Independent of the MSE direction, the ACF in Figure "
        "A3 still endorses the paper's primary explanatory contribution: "
        "ML's advantage on realised variance is a long-memory advantage.",
    ])

    add_page_break(DOC)

    # ===== APPENDIX =====
    add_heading(DOC, "Appendix", level=1)

    # --- Table A1: Descriptive ---
    add_heading(DOC, "A. Tables", level=2)
    df1 = pd.read_csv(ROOT / "tab1_descriptive.csv")
    make_dataframe_table(DOC, df1, font_size=9)
    add_caption(DOC, "Table A1.",
                "Descriptive statistics of daily realised variance "
                "(5-minute subsampled). RV is reported scaled by 10⁴; "
                "annualised σ = √(252·RV).")

    # --- Table A2: Main results ---
    df2 = pd.read_csv(ROOT / "tab2_main_results.csv")
    make_dataframe_table(DOC, df2, font_size=9)
    add_caption(DOC, "Table A2.",
                "Out-of-sample MSE relative to HAR. Lower is better; "
                "1.000 = HAR baseline. Stars indicate Diebold–Mariano "
                "rejection of equal predictive accuracy in favour of the "
                "row model, with Harvey–Leybourne–Newbold small-sample "
                "correction (*: 10 %, **: 5 %, ***: 1 %). EN = Elastic "
                "Net; RF = Random Forest; NN = Neural Network (4, 2).")

    # --- Table A3: QLIKE robustness ---
    df3 = pd.read_csv(ROOT / "tab3_qlike.csv")
    make_dataframe_table(DOC, df3, font_size=9)
    add_caption(DOC, "Table A3.",
                "QLIKE robustness check (Patton, 2011). "
                "Cross-section mean of out-of-sample QLIKE relative to "
                "HAR; lower is better. A '*' indicates DM rejection on "
                "the QLIKE differential in at least 2 of 3 stocks at the "
                "10 % level. RF's monthly deterioration shrinks from "
                "2.142 (MSE) to 1.296 (QLIKE); HARQ's monthly 1.190 (MSE) "
                "collapses to 0.959 — both consistent with proxy-noise "
                "amplification of absolute errors that QLIKE discounts.")

    add_page_break(DOC)

    # --- Figures ---
    add_heading(DOC, "B. Figures", level=2)

    add_image(DOC, "fig1_rv_overview.png", width_in=6.4,
              label="Figure A1.",
              caption="Daily 5-minute realised volatility (annualised, %) "
                      "for AAPL, AMZN and JPM, 2016–2024. COVID-19 "
                      "(March 2020) and the August 2024 carry-trade "
                      "unwind are clearly visible. The 70/10/20 split "
                      "places the test window at March 2023 – December "
                      "2024, an unusually quiescent period.")

    add_image(DOC, "fig3_relative_mse.png", width_in=6.4,
              label="Figure A2.",
              caption="Out-of-sample MSE relative to HAR for each "
                      "model–stock–horizon. The dashed line at 1.0 is "
                      "the HAR benchmark; values below 1.0 indicate "
                      "improvement. LogHAR is the strongest HAR variant "
                      "at every horizon; the Random Forest deteriorates "
                      "sharply at h = 22, especially on AMZN.")

    add_image(DOC, "fig2_acf_persistence.png", width_in=5.2,
              label="Figure A3.",
              caption="Reproduction of CSV's Figure 8. In-sample "
                      "autocorrelation of fitted realised variance at "
                      "the short (h = 1) and long (h = 22) horizons for "
                      "HAR, LogHAR, Random Forest and Neural Network. "
                      "At h = 22 the ML models retain significant "
                      "autocorrelation beyond lag 60 while HAR has "
                      "decayed close to zero — independent visual "
                      "confirmation of CSV's long-memory mechanism.")

    add_page_break(DOC)

    add_image(DOC, "fig4_forecasts.png", width_in=6.4,
              label="Figure A4.",
              caption="AAPL one-day-ahead forecasts versus realised "
                      "volatility over the test window. HAR, LogHAR and "
                      "NN are visually indistinguishable in the centre "
                      "of the distribution; differences emerge at "
                      "volatility spikes (August 2024).")

    add_image(DOC, "fig5_decile_mse.png", width_in=6.4,
              label="Figure A5.",
              caption="Reproduction of CSV's Figure 5. Forecast MSE "
                      "relative to HAR by decile of realised RV in the "
                      "test set, h = 1. LogHAR is uniformly better than "
                      "HAR for AAPL and JPM; RF and NN are worse in the "
                      "middle deciles for AMZN but converge to HAR in "
                      "the top decile, where LogHAR's outlier-down-"
                      "weighting dominates.")

    add_image(DOC, "fig6_rolling_robust.png", width_in=5.5,
              label="Figure A6.",
              caption="Robustness of the long-horizon Random-Forest "
                      "deterioration. Re-estimating HAR and RF every "
                      "22 days with an expanding window leaves the "
                      "RF/HAR ratio essentially unchanged on JPM, "
                      "marginally lower on AAPL and AMZN. The "
                      "long-horizon divergence is not an artefact of "
                      "static estimation.")

    add_page_break(DOC)

    # --- References ---
    add_heading(DOC, "References", level=1)

    refs = [
        "Andersen, T. G., & Bollerslev, T. (1998). Answering the skeptics: "
        "Yes, standard volatility models do provide accurate forecasts. "
        "International Economic Review, 39(4), 885–905.",
        "Apley, D. W., & Zhu, J. (2020). Visualizing the effects of "
        "predictor variables in black box supervised learning models. "
        "Journal of the Royal Statistical Society: Series B, 82(4), "
        "1059–1086.",
        "Bollerslev, T., Patton, A. J., & Quaedvlieg, R. (2016). "
        "Exploiting the errors: A simple approach for improved volatility "
        "forecasting. Journal of Econometrics, 192(1), 1–18.",
        "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5–32.",
        "Christensen, K., Siggaard, M., & Veliyev, B. (2023). A machine "
        "learning approach to volatility forecasting. Journal of "
        "Financial Econometrics, 21(5), 1680–1727.",
        "Corsi, F. (2009). A simple approximate long-memory model of "
        "realized volatility. Journal of Financial Econometrics, 7(2), "
        "174–196.",
        "Diebold, F. X., & Mariano, R. S. (1995). Comparing predictive "
        "accuracy. Journal of Business & Economic Statistics, 13(3), "
        "253–263.",
        "Harvey, D., Leybourne, S., & Newbold, P. (1997). Testing the "
        "equality of prediction mean squared errors. International "
        "Journal of Forecasting, 13(2), 281–291.",
        "Kingma, D. P., & Ba, J. (2014). Adam: A method for stochastic "
        "optimization. arXiv:1412.6980.",
        "Patton, A. J. (2011). Volatility forecast comparison using "
        "imperfect volatility proxies. Journal of Econometrics, 160(1), "
        "246–256.",
        "Patton, A. J., & Sheppard, K. (2015). Good volatility, bad "
        "volatility: Signed jumps and the persistence of volatility. "
        "Review of Economics and Statistics, 97(3), 683–697.",
        "Zou, H., & Hastie, T. (2005). Regularization and variable "
        "selection via the elastic net. Journal of the Royal Statistical "
        "Society: Series B, 67(2), 301–320.",
    ]
    for ref in refs:
        p = DOC.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(ref)
        run.font.size = Pt(10)

    # ===== SAVE =====
    out_path = ROOT / "Volatility_Forecasting_Replication_Report.docx"
    DOC.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


if __name__ == "__main__":
    build_report()
